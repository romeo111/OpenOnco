#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Generate docs/moz/draft_nakaz_moz_dod4_2026-06-01.docx
from KB data using python-docx.

A4, Times New Roman, proper heading styles, 8-column МОЗ table.
"""

import sys, os, yaml, re

# ─── ensure python-docx available ────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("python-docx not installed: pip install python-docx")

# reuse logic from gen_nakaz_dod4.py ──────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, SCRIPT_DIR)

# Import the generator logic directly
from gen_nakaz_dod4 import (
    KB_BASE, load_tests, load_biomarkers, load_diseases,
    load_indications_by_disease, classify_test, nice_test_name,
    parse_icd10_sort_key, is_nccn_or_esmo,
)

OUT_FILE = "C:/Users/805/cancer-autoresearch/docs/moz/draft_nakaz_moz_dod4_2026-06-01.docx"

# Colours
HDR_FILL  = "BDD7EE"  # light blue
ALT_FILL  = "F2F2F2"  # light grey
WHITE     = "FFFFFF"


# ─── helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def cell_font(cell, size_pt: int = 9, bold: bool = False, color_hex: str | None = None):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name  = "Times New Roman"
            run.font.size  = Pt(size_pt)
            run.font.bold  = bold
            if color_hex:
                run.font.color.rgb = RGBColor.from_string(color_hex)


def add_header_row(table, headers: list[str]):
    hdr_row = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = text
        set_cell_bg(cell, HDR_FILL)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name  = "Times New Roman"
                run.font.size  = Pt(9)
                run.font.bold  = True


def add_data_row(table, values: list[str], alt: bool):
    row = table.add_row()
    fill = ALT_FILL if alt else WHITE
    for i, text in enumerate(values):
        cell = row.cells[i]
        cell.text = str(text)
        set_cell_bg(cell, fill)
        cell_font(cell, size_pt=9)


def set_col_width(table, col_widths_cm: list[float]):
    for row in table.rows:
        for i, w in enumerate(col_widths_cm):
            row.cells[i].width = Cm(w)


def heading(doc, text: str, level: int):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
    return p


def body_para(doc, text: str, bold: bool = False, center: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.bold = bold
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


# ─── build rows from KB ───────────────────────────────────────────────────────

def build_rows():
    print("Loading KB data...", file=sys.stderr)
    tests_dir = os.path.join(KB_BASE, "tests")
    bm_dir    = os.path.join(KB_BASE, "biomarkers")
    dis_dir   = os.path.join(KB_BASE, "diseases")
    ind_dir   = os.path.join(KB_BASE, "indications")

    test_map   = load_tests(tests_dir)
    bm_map     = load_biomarkers(bm_dir)
    diseases   = load_diseases(dis_dir)
    ind_by_dis = load_indications_by_disease(ind_dir)

    print(f"  {len(diseases)} diseases, {len(test_map)} tests, "
          f"{len(bm_map)} biomarkers, "
          f"{sum(len(v) for v in ind_by_dis.values())} indications",
          file=sys.stderr)

    diseases.sort(key=lambda d: (
        parse_icd10_sort_key(d.get("codes", {}).get("icd_10")),
        d.get("id", "")
    ))

    rows = []
    all_sources = set()

    for idx, dis in enumerate(diseases, start=1):
        did   = dis.get("id", "")
        codes = dis.get("codes", {})
        icd10  = str(codes.get("icd_10", "") or "—")
        icd_o  = str(codes.get("icdo4_code") or codes.get("icd_o_3_morphology", "") or "—")
        names  = dis.get("names", {})
        name_ua = names.get("ukrainian") or names.get("preferred") or did

        indications  = ind_by_dis.get(did, [])
        req_tests    = set()
        bm_req       = {}
        sources      = set()

        for ind in indications:
            for t in (ind.get("required_tests") or []):
                req_tests.add(t)
            for bmr in (ind.get("applicable_to", {}).get("biomarker_requirements_required") or []):
                bid = bmr.get("biomarker_id")
                if bid:
                    bm_req[bid] = bmr.get("value_constraint", "")
            for src in (ind.get("sources") or []):
                sid = src.get("source_id") if isinstance(src, dict) else str(src)
                if sid and is_nccn_or_esmo(sid):
                    sources.add(sid)

        all_sources.update(sources)

        clinical_tests = []
        ihc_tests      = []
        molecular_tests= []

        for t in sorted(req_tests):
            cat  = classify_test(t)
            name = nice_test_name(t, test_map)
            if cat in ("clinical", "serology"):
                clinical_tests.append(name)
            elif cat == "ihc":
                ihc_tests.append(name)
            elif cat == "molecular":
                molecular_tests.append(name)

        bm_parts = []
        for bid in sorted(bm_req.keys()):
            bm      = bm_map.get(bid, {})
            bm_names = bm.get("names", {})
            fallback = bid.replace("BIO-","").replace("-"," ").title()
            bm_ua   = bm_names.get("ukrainian") or bm_names.get("preferred") or fallback
            bm_parts.append(bm_ua)

        placeholder = "— (підлягає визначенню)"
        has_ind = bool(indications)

        def col(lst):
            return "; ".join(lst) if lst else ""

        rows.append({
            "num":       idx,
            "icd_o":     icd_o,
            "icd10":     icd10,
            "name":      name_ua,
            "clinical":  col(clinical_tests) if has_ind else placeholder,
            "ihc":       col(ihc_tests)       if has_ind else placeholder,
            "molecular": col(molecular_tests) if has_ind else placeholder,
            "biomarkers":col(bm_parts),
            "has_ind":   has_ind,
        })

    no_ind = [r for r in rows if not r["has_ind"]]
    if no_ind:
        print(f"\nDiseases with NO indications ({len(no_ind)}):")
        for r in no_ind:
            print(f"  - {r['name']} ({r['icd10']})", file=sys.stderr)
    else:
        print("\nAll diseases have at least one indication.", file=sys.stderr)

    return rows, sorted(all_sources)


# ─── main ─────────────────────────────────────────────────────────────────────

def main():
    rows, sources = build_rows()

    doc = Document()

    # Page setup: A4
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.0)
    section.top_margin  = section.bottom_margin = Cm(2.0)

    # ── Document header ────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ПРОЄКТ")
    run.font.name = "Times New Roman"; run.font.size = Pt(14); run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("МІНІСТЕРСТВО ОХОРОНИ ЗДОРОВ'Я УКРАЇНИ")
    run.font.name = "Times New Roman"; run.font.size = Pt(14); run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("НАКАЗ")
    run.font.name = "Times New Roman"; run.font.size = Pt(14); run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("від 01 червня 2026 року № ___")
    run.font.name = "Times New Roman"; run.font.size = Pt(12); run.font.bold = True

    heading(doc, "Про доповнення наказу Міністерства охорони здоров'я України від 07 квітня 2026 року № 473", level=3)

    body_para(doc,
        "Відповідно до статей 14, 34 Закону України «Про основи законодавства України про охорону здоров'я», "
        "підпункту 3 пункту 4 Положення про Міністерство охорони здоров'я України, затвердженого постановою "
        "Кабінету Міністрів України від 25 березня 2015 року № 267, з метою забезпечення стандартизованого "
        "підходу до молекулярно-генетичної та імуногістохімічної діагностики злоякісних новоутворень "
        "відповідно до сучасних міжнародних стандартів (NCCN, ESMO, CAP)"
    )

    body_para(doc, "НАКАЗУЮ:", bold=True)

    body_para(doc,
        "1. Доповнити наказ Міністерства охорони здоров'я України від 07 квітня 2026 року № 473 "
        "«Про затвердження довідників для кодування та формування онкологічного діагнозу» додатком 4 "
        "«Перелік рекомендованих молекулярно-генетичних та імуногістохімічних досліджень при встановленні "
        "онкологічного діагнозу», що додається."
    )
    body_para(doc,
        "2. Керівникам закладів охорони здоров'я, що надають спеціалізовану онкологічну допомогу, "
        "забезпечити виконання вимог додатку 4 при формуванні онкологічного діагнозу."
    )
    body_para(doc,
        "3. Контроль за виконанням цього наказу покласти на заступника Міністра охорони здоров'я України "
        "відповідно до розподілу обов'язків."
    )
    body_para(doc, "Міністр охорони здоров'я України                                            [підпис]")

    doc.add_paragraph("─" * 80)

    # ── Annex header ──────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(
        "ЗАТВЕРДЖЕНО\nНаказ Міністерства охорони здоров'я України\n"
        "від 01 червня 2026 року № ___"
    )
    run.font.name = "Times New Roman"; run.font.size = Pt(12); run.font.bold = True

    heading(doc, "Додаток 4 до наказу МОЗ від 07.04.2026 № 473", level=2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "ПЕРЕЛІК\nрекомендованих молекулярно-генетичних та імуногістохімічних досліджень\n"
        "при встановленні онкологічного діагнозу"
    )
    run.font.name = "Times New Roman"; run.font.size = Pt(13); run.font.bold = True

    doc.add_paragraph("")

    # ── Table ─────────────────────────────────────────────────────────────────
    headers = [
        "№", "ICD-O-4", "ICD-10", "Нозологія",
        "Загальноклінічні та серологічні",
        "ІГХ та морфологія",
        "Молекулярна генетика",
        "Ключові молекулярні маркери",
    ]
    # Approximate column widths (total ≈ 17 cm for A4 with 2 cm margins each side)
    col_widths = [0.8, 1.5, 1.5, 3.5, 3.0, 2.5, 2.5, 2.0]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    add_header_row(table, headers)

    for i, row in enumerate(rows):
        values = [
            str(row["num"]),
            row["icd_o"],
            row["icd10"],
            row["name"],
            row["clinical"],
            row["ihc"],
            row["molecular"],
            row["biomarkers"],
        ]
        add_data_row(table, values, alt=(i % 2 == 1))

    set_col_width(table, col_widths)

    doc.add_paragraph("")

    # ── Sources ───────────────────────────────────────────────────────────────
    if sources:
        heading(doc, "Джерела", level=2)
        body_para(doc,
            "Перелік сформовано на підставі наступних міжнародних клінічних настанов, "
            "на які посилаються індикації бази знань OpenOnco:"
        )
        for src in sources:
            doc.add_paragraph(f"• {src}", style="List Bullet")

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(
        "Підготовлено: OpenOnco (openonco.info). "
        "Джерела: NCCN Clinical Practice Guidelines in Oncology, "
        "ESMO Clinical Practice Guidelines, CAP Cancer Protocols. "
        "Версія 1.0, травень 2026."
    )
    run.font.name = "Times New Roman"; run.font.size = Pt(10); run.font.italic = True

    os.makedirs(os.path.dirname(OUT_FILE), exist_ok=True)
    doc.save(OUT_FILE)
    size = os.path.getsize(OUT_FILE)
    print(f"\nSaved: {OUT_FILE}  ({size:,} bytes, {size/1024:.1f} KB)", file=sys.stderr)
    print(f"Saved: {OUT_FILE}  ({size:,} bytes, {size/1024:.1f} KB)")


if __name__ == "__main__":
    main()
